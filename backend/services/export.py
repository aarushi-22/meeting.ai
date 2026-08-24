"""
Export helpers: Word doc (transcript + summary) and .ics (action item deadlines).

Both are generated on-demand from what's already in SQLite — nothing extra
persisted for exports, no export-specific storage.
"""

import io
from datetime import datetime, date
from docx import Document
from docx.shared import Pt
from icalendar import Calendar, Event


def build_meeting_docx(meeting: dict) -> io.BytesIO:
    doc = Document()

    doc.add_heading(meeting["filename"], level=1)
    meta = doc.add_paragraph()
    meta.add_run(f"Processed: {meeting['created_at']}").italic = True

    doc.add_heading("Summary", level=2)
    doc.add_paragraph(meeting.get("summary") or "No summary available.")

    doc.add_heading("Action Items", level=2)
    items = meeting.get("action_items") or []
    if not items:
        doc.add_paragraph("No action items extracted.")
    else:
        for item in items:
            p = doc.add_paragraph(style="List Bullet")
            owner = item.get("owner") or "Unassigned"
            deadline = item.get("deadline") or "No deadline"
            p.add_run(f"{item['task']} ").bold = True
            p.add_run(f"(Owner: {owner} · Due: {deadline})").italic = True

    doc.add_heading("Full Transcript", level=2)
    transcript_para = doc.add_paragraph(meeting.get("transcript") or "")
    for run in transcript_para.runs:
        run.font.size = Pt(10)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def build_action_items_ics(meeting: dict) -> io.BytesIO:
    """One .ics file containing one all-day VEVENT per action item that has a deadline.
    Items with no deadline are skipped — nothing sensible to put on a calendar.
    """
    cal = Calendar()
    cal.add("prodid", "-//Meeting Summarizer//local//")
    cal.add("version", "2.0")

    items = meeting.get("action_items") or []
    added = 0
    for item in items:
        if not item.get("deadline"):
            continue
        try:
            due = datetime.strptime(item["deadline"], "%Y-%m-%d").date()
        except ValueError:
            continue

        event = Event()
        event.add("summary", item["task"])
        owner = item.get("owner") or "Unassigned"
        event.add("description", f"From meeting: {meeting['filename']}\nOwner: {owner}")
        event.add("dtstart", due)
        event.add("dtend", due)
        event.add("dtstamp", datetime.utcnow())
        event["uid"] = f"{meeting['id']}-{item['id']}@meeting-summarizer.local"
        cal.add_component(event)
        added += 1

    buf = io.BytesIO()
    buf.write(cal.to_ical())
    buf.seek(0)
    return buf, added
